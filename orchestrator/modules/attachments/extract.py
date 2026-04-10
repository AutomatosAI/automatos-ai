"""
Text extraction for document attachments (PRD-127)

Extracts plain text from:
- PDF (pdfplumber)
- DOCX (python-docx)
- XLSX (openpyxl)
- CSV, TXT, code files (direct read)

This module has NO imports from modules/rag/ or DocumentManager.
It's a standalone extraction utility for ephemeral attachments.
"""

from __future__ import annotations

import csv
import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text(
    content: bytes,
    mime: str,
    filename: str,
    max_chars: Optional[int] = None,
) -> str:
    """
    Extract text from document bytes.

    Args:
        content: Raw file bytes
        mime: MIME type (determines extraction method)
        filename: Original filename (fallback for type detection)
        max_chars: Optional character limit

    Returns:
        Extracted text (may be truncated if max_chars specified)
    """
    try:
        if mime == "application/pdf":
            text = _extract_pdf(content)
        elif mime in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            text = _extract_docx(content)
        elif mime in (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
        ):
            text = _extract_xlsx(content)
        elif mime == "text/csv":
            text = _extract_csv(content)
        elif mime.startswith("text/") or mime in (
            "application/json",
            "application/xml",
            "application/x-yaml",
            "application/x-sh",
        ):
            text = _extract_text_file(content)
        else:
            # Unknown format — try as text, fall back to placeholder
            try:
                text = content.decode("utf-8", errors="replace")
            except Exception:
                text = f"[Binary file: {filename}]"

        # Truncate if needed
        if max_chars and len(text) > max_chars:
            text = text[:max_chars] + f"\n\n[... truncated at {max_chars} chars]"

        return text.strip()

    except Exception as e:
        logger.error("Text extraction failed for %s (%s): %s", filename, mime, e)
        return f"[Extraction failed for {filename}: {e}]"


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber not installed, cannot extract PDF")
        return "[PDF extraction requires pdfplumber]"

    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"--- Page {i + 1} ---\n{page_text}")

    if not text_parts:
        # PDF might be image-only — note this for the user
        return "[PDF contains no extractable text — may be image-only or scanned]"

    return "\n\n".join(text_parts)


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed, cannot extract DOCX")
        return "[DOCX extraction requires python-docx]"

    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


def _extract_xlsx(content: bytes) -> str:
    """Extract text from XLSX using openpyxl."""
    try:
        import openpyxl
    except ImportError:
        logger.warning("openpyxl not installed, cannot extract XLSX")
        return "[XLSX extraction requires openpyxl]"

    wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
    text_parts = []

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        rows = []
        for row in sheet.iter_rows(values_only=True):
            # Filter out empty cells, convert to strings
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(" | ".join(cells))

        if rows:
            text_parts.append(f"### Sheet: {sheet_name}\n" + "\n".join(rows))

    return "\n\n".join(text_parts) if text_parts else "[Empty spreadsheet]"


def _extract_csv(content: bytes) -> str:
    """Extract text from CSV."""
    try:
        text = content.decode("utf-8", errors="replace")
    except Exception:
        text = content.decode("latin-1", errors="replace")

    # Parse and format as table
    try:
        reader = csv.reader(io.StringIO(text))
        rows = [" | ".join(row) for row in reader if any(row)]
        return "\n".join(rows)
    except Exception:
        # Fallback to raw text
        return text


def _extract_text_file(content: bytes) -> str:
    """Extract text from plain text / code files."""
    # Try UTF-8 first, then fall back to latin-1
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1", errors="replace")
