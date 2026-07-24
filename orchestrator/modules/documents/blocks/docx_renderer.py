"""Block document → DOCX renderer (PRD-167 S2, Q71).

Compiles a :class:`BlockDocument` directly into a ``python-docx`` Document — the *same
block tree* that drives the PDF path. This removes the previous requirement that DOCX
generation needs a pre-uploaded ``.docx`` template file (Q71): a block template renders
to both PDF and DOCX from one source.

python-docx is pure-Python (no system libraries), so this path needs no Docker / native
deps to build or unit-test.

Brand kit (PRD-167 S4): heading colour comes from ``brand.primary_color`` and the body
font from ``brand.font_family`` — no hardcoded Automatos styling.
"""

from __future__ import annotations

import ipaddress
import logging
import os
import socket
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional
from urllib.parse import urlparse
from urllib.request import HTTPRedirectHandler, Request, build_opener

from .schema import BlockDocument

logger = logging.getLogger(__name__)

_MAX_IMAGE_BYTES = 10 * 1024 * 1024  # 10 MB cap on fetched images
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}


class _NoRedirect(HTTPRedirectHandler):
    """Refuse to follow redirects — a public URL must not 30x into a private host
    (SSRF). Returning None makes urllib raise instead of following."""

    def redirect_request(self, req, fp, code, msg, headers, newurl):  # noqa: D401
        return None


@dataclass
class RenderedDocx:
    document: object  # docx.document.Document (typed loosely to avoid import at module load)
    unresolved: List[str] = field(default_factory=list)


def _hex_to_rgb(value: Optional[str]):
    from docx.shared import RGBColor

    if not value:
        return None
    v = value.lstrip("#")
    if len(v) == 3:
        v = "".join(c * 2 for c in v)
    try:
        return RGBColor(int(v[0:2], 16), int(v[2:4], 16), int(v[4:6], 16))
    except (ValueError, IndexError):
        return None


def _safe_local_image(src: str) -> Optional[BytesIO]:
    """Read a local/upload image, confined under the workspace document-storage root.

    Rejects path traversal (``../``, absolute paths escaping the root) and non-image
    extensions, so a template can't coerce an arbitrary-file read (e.g. ``/etc/passwd``)."""
    if os.path.splitext(src)[1].lower() not in _IMAGE_EXTS:
        return None
    try:
        from config import config

        root = Path(config.DOCUMENT_STORAGE_DIR).resolve()
    except Exception:  # noqa: BLE001 — config unavailable
        return None
    candidate = (root / src.lstrip("/\\")).resolve()
    if root != candidate and root not in candidate.parents:
        logger.warning("[DocxRender] refusing image path outside storage root: %r", src)
        return None
    try:
        with open(candidate, "rb") as fh:
            return BytesIO(fh.read(_MAX_IMAGE_BYTES + 1))
    except OSError:
        return None


def _safe_image_bytes(src: str) -> Optional[BytesIO]:
    """Best-effort, SSRF-guarded image fetch for DOCX embedding.

    Local/upload paths are read only from within the document-storage root
    (:func:`_safe_local_image`); http(s) URLs are fetched only when every resolved
    address is public AND redirects are refused (a 30x must not pivot into a private
    host — mirrors and tightens the PRD-156 S4 WeasyPrint SSRF posture). Any failure
    returns ``None`` (the caller falls back to alt text)."""
    if not src:
        return None
    if not src.startswith(("http://", "https://")):
        return _safe_local_image(src)
    parsed = urlparse(src)
    host = parsed.hostname or ""
    try:
        infos = socket.getaddrinfo(host, None)
    except socket.gaierror:
        return None
    for info in infos:
        if not ipaddress.ip_address(info[4][0]).is_global:
            logger.warning("[DocxRender] refusing non-public image host %r", host)
            return None
    opener = build_opener(_NoRedirect)
    try:
        with opener.open(Request(src, headers={"User-Agent": "Automatos-DocGen"}), timeout=5) as resp:  # noqa: S310 — host validated + redirects refused
            if getattr(resp, "status", 200) in (301, 302, 303, 307, 308):
                return None
            data = resp.read(_MAX_IMAGE_BYTES + 1)
        if len(data) > _MAX_IMAGE_BYTES:
            return None
        return BytesIO(data)
    except Exception:  # noqa: BLE001
        return None


def _resolve_var(path: str, fallback, values: Dict[str, str], unresolved: List[str]) -> str:
    if path in values:
        return values[path]
    if fallback is not None:
        return fallback
    unresolved.append(path)
    return f"[[{path}]]"


def _add_inline(paragraph, content: list, values: Dict[str, str], unresolved: List[str], font: Optional[str]):
    for run_spec in content:
        if run_spec.type == "text":
            run = paragraph.add_run(run_spec.text)
            run.bold = "bold" in run_spec.marks
            run.italic = "italic" in run_spec.marks
            run.underline = "underline" in run_spec.marks
            if font:
                run.font.name = font
        elif run_spec.type == "variable":
            text = _resolve_var(run_spec.path, run_spec.fallback, values, unresolved)
            run = paragraph.add_run(text)
            if font:
                run.font.name = font


def _add_block(doc, block, values, brand_kit, unresolved, *, primary_rgb, font):
    from docx.shared import Mm

    kind = block.type
    if kind == "heading":
        p = doc.add_heading(level=min(block.level, 9))
        p.clear()
        _add_inline(p, block.content, values, unresolved, font)
        if primary_rgb is not None:
            for run in p.runs:
                run.font.color.rgb = primary_rgb
    elif kind == "text":
        p = doc.add_paragraph()
        _add_inline(p, block.content, values, unresolved, font)
    elif kind == "table":
        n_rows = len(block.rows)
        n_cols = max((len(r) for r in block.rows), default=0)
        if n_rows and n_cols:
            table = doc.add_table(rows=n_rows, cols=n_cols)
            table.style = "Light Grid Accent 1"
            for r_idx, row in enumerate(block.rows):
                for c_idx, cell in enumerate(row):
                    para = table.cell(r_idx, c_idx).paragraphs[0]
                    _add_inline(para, cell, values, unresolved, font)
                    if block.header and r_idx == 0:
                        for run in para.runs:
                            run.bold = True
    elif kind == "image":
        src = (brand_kit or {}).get("logo_url", "") if block.source == "brand_logo" else (block.src or "")
        if block.source == "brand_logo" and not src:
            unresolved.append("brand.logo_url")
            return
        stream = _safe_image_bytes(src)
        if stream is not None:
            try:
                width = Mm(block.width_mm) if block.width_mm else Mm(60)
                doc.add_picture(stream, width=width)
            except Exception:  # noqa: BLE001 — unreadable image format
                doc.add_paragraph(block.alt or "")
        else:
            doc.add_paragraph(block.alt or "")
    elif kind == "variable":
        p = doc.add_paragraph()
        p.add_run(_resolve_var(block.path, block.fallback, values, unresolved))
    elif kind == "page_break":
        doc.add_page_break()
    elif kind == "section":
        if block.title:
            h = doc.add_heading(block.title, level=2)
            if primary_rgb is not None:
                for run in h.runs:
                    run.font.color.rgb = primary_rgb
        for child in block.children:
            _add_block(doc, child, values, brand_kit, unresolved, primary_rgb=primary_rgb, font=font)


def render_document_docx(doc_model: BlockDocument, values: Dict[str, str], brand_kit: Dict) -> RenderedDocx:
    """Render a block document to a python-docx Document + unresolved-path list."""
    from docx import Document

    document = Document()
    bk = brand_kit or {}
    primary_rgb = _hex_to_rgb(bk.get("primary_color"))
    # font_family may be a CSS stack ("Inter, 'Segoe UI', ..."); take the first family.
    font_stack = (bk.get("font_family") or "").split(",")[0].strip().strip("'\"") or None

    unresolved: List[str] = []
    for block in doc_model.blocks:
        _add_block(document, block, values, bk, unresolved, primary_rgb=primary_rgb, font=font_stack)

    seen: Dict[str, None] = {}
    for path in unresolved:
        seen.setdefault(path, None)
    return RenderedDocx(document=document, unresolved=list(seen))


__all__ = ["RenderedDocx", "render_document_docx"]
